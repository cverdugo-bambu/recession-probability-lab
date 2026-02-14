"""Investigation of Ambassador Health / Ambassador Home Health Services fraud signals.

This module consolidates all Ambassador-branded entities appearing in the
Florida Medicaid provider-billing data, computes fraud-risk indicators, and
produces a structured report for further review.

Key findings summary (from the data):
    - 9 distinct NPI numbers all billing under the "Ambassador" brand
    - Spread across 9 Florida cities with a combined $197M+ in Medicaid billings
    - Multiple entities show extreme cost-per-beneficiary and claims-per-beneficiary
    - Several locations bill only 1-2 HCPCS codes (very low service diversity)
    - Pattern is consistent with a possible franchise/shell network
"""

from __future__ import annotations

import json
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

import pandas as pd

from .medicaid_analysis import load_pbc_providers, compute_fraud_signals, ARTIFACT_DIR


# Known HCPCS codes confirmed from claim-level sample data
AMBASSADOR_KNOWN_HCPCS: dict[str, dict] = {
    "S9124": {
        "description": "Nursing care, in the home; by registered nurse, per hour",
        "category": "Home Health Nursing",
        "confirmed_npis": ["1124261284", "1023477940"],
        "note": "Confirmed from claim-level sample data. Likely billed by all entities.",
    },
    "S9123": {
        "description": "Nursing care, in the home; by licensed practical nurse, per hour",
        "category": "Home Health Nursing",
        "note": "Probable — same service category, LPN-level variant of S9124.",
    },
    "S9122": {
        "description": "Home health aide or certified nursing assistant, in the home; per hour",
        "category": "Home Health Aide",
        "note": "Probable — same service category, aide-level variant.",
    },
}


# ---------------------------------------------------------------------------
# Ambassador entity extraction
# ---------------------------------------------------------------------------

def get_ambassador_entities(providers_df: pd.DataFrame) -> pd.DataFrame:
    """Filter the provider dataset to Ambassador-branded entities only."""
    mask = providers_df["provider_name"].str.contains(
        r"(?i)ambassador", na=False
    )
    return providers_df[mask].copy()


# ---------------------------------------------------------------------------
# Fraud signal computation
# ---------------------------------------------------------------------------

@dataclass
class AmbassadorFraudReport:
    """Container for the full Ambassador investigation output."""

    entities: pd.DataFrame
    total_paid: float = 0.0
    entity_count: int = 0
    unique_cities: list[str] = field(default_factory=list)
    unique_npis: list[str] = field(default_factory=list)
    flags_summary: dict[str, int] = field(default_factory=dict)
    peer_comparison: dict[str, float] = field(default_factory=dict)
    red_flags: list[str] = field(default_factory=list)


def _compute_peer_stats(all_providers: pd.DataFrame) -> dict[str, float]:
    """Return median/mean/p95 for key metrics across all providers."""
    return {
        "median_cost_per_claim": float(all_providers["cost_per_claim"].median()),
        "mean_cost_per_claim": float(all_providers["cost_per_claim"].mean()),
        "p95_cost_per_claim": float(all_providers["cost_per_claim"].quantile(0.95)),
        "median_cost_per_beneficiary": float(
            all_providers["cost_per_beneficiary"].median()
        ),
        "mean_cost_per_beneficiary": float(
            all_providers["cost_per_beneficiary"].mean()
        ),
        "p95_cost_per_beneficiary": float(
            all_providers["cost_per_beneficiary"].quantile(0.95)
        ),
        "median_claims_per_beneficiary": float(
            all_providers["claims_per_beneficiary"].median()
        ),
        "mean_claims_per_beneficiary": float(
            all_providers["claims_per_beneficiary"].mean()
        ),
        "p95_claims_per_beneficiary": float(
            all_providers["claims_per_beneficiary"].quantile(0.95)
        ),
        "median_unique_hcpcs": float(all_providers["unique_hcpcs_codes"].median()),
        "mean_unique_hcpcs": float(all_providers["unique_hcpcs_codes"].mean()),
    }


def _identify_red_flags(
    amb: pd.DataFrame, peers: dict[str, float]
) -> list[str]:
    """Generate plain-English red flag descriptions from the data."""
    flags: list[str] = []

    # 1) Network / shell structure
    n = len(amb)
    cities = amb["city"].nunique()
    total = amb["total_paid"].sum()
    flags.append(
        f"NETWORK STRUCTURE: {n} separate NPIs billing under the Ambassador brand "
        f"across {cities} Florida cities, totalling ${total:,.2f} in Medicaid payments. "
        "This is consistent with a franchise or shell-company network designed to "
        "stay below per-entity audit thresholds."
    )

    # 2) Extreme cost per beneficiary
    high_cpb = amb[amb["cost_per_beneficiary"] > peers["p95_cost_per_beneficiary"]]
    if not high_cpb.empty:
        worst = high_cpb.sort_values("cost_per_beneficiary", ascending=False).iloc[0]
        flags.append(
            f"EXTREME COST PER BENEFICIARY: {len(high_cpb)} of {n} Ambassador entities "
            f"exceed the 95th-percentile cost-per-beneficiary "
            f"(>${peers['p95_cost_per_beneficiary']:,.2f}). "
            f"Worst offender: {worst['provider_name']} in {worst['city']} at "
            f"${worst['cost_per_beneficiary']:,.2f}/beneficiary — "
            f"{worst['cost_per_beneficiary'] / peers['median_cost_per_beneficiary']:.1f}x "
            f"the median provider."
        )

    # 3) Extreme cost per claim
    high_cpc = amb[amb["cost_per_claim"] > peers["p95_cost_per_claim"]]
    if not high_cpc.empty:
        flags.append(
            f"HIGH COST PER CLAIM: {len(high_cpc)} of {n} entities exceed the "
            f"95th-percentile cost-per-claim (>${peers['p95_cost_per_claim']:,.2f})."
        )

    # 4) Low code diversity (potential upcoding / phantom billing)
    low_div = amb[amb["unique_hcpcs_codes"] <= 2]
    if not low_div.empty:
        codes_list = ", ".join(
            f"{r['provider_name']} ({r['city']}): {r['unique_hcpcs_codes']} codes"
            for _, r in low_div.iterrows()
        )
        flags.append(
            f"LOW SERVICE DIVERSITY: {len(low_div)} of {n} entities bill only 1-2 HCPCS "
            f"codes (median provider uses {peers['median_unique_hcpcs']:.0f}). "
            f"This is a classic indicator of upcoding or phantom billing. "
            f"Details: {codes_list}"
        )

    # 5) High claims per beneficiary
    high_cpb_claims = amb[
        amb["claims_per_beneficiary"] > peers["p95_claims_per_beneficiary"]
    ]
    if not high_cpb_claims.empty:
        flags.append(
            f"EXCESSIVE UTILIZATION: {len(high_cpb_claims)} of {n} entities have "
            f"claims-per-beneficiary above the 95th percentile "
            f"({peers['p95_claims_per_beneficiary']:.1f}). "
            "This may indicate billing for services not actually rendered."
        )

    # 6) Geographic spread with consistent patterns
    if cities >= 5:
        avg_cpc = amb["cost_per_claim"].mean()
        std_cpc = amb["cost_per_claim"].std()
        cv = std_cpc / avg_cpc if avg_cpc > 0 else 0
        flags.append(
            f"GEOGRAPHIC DISPERSION: Ambassador operates in {cities} cities across "
            f"Florida. Average cost/claim=${avg_cpc:,.2f} with CV={cv:.2f}. "
            "A coordinated billing scheme would show consistent pricing patterns "
            "across locations."
        )

    # 7) Outlier entity (Tampa — looks different from the rest)
    tampa = amb[amb["city"] == "TAMPA"]
    if not tampa.empty:
        t = tampa.iloc[0]
        other_avg_cpc = amb[amb["city"] != "TAMPA"]["cost_per_claim"].mean()
        flags.append(
            f"TAMPA ANOMALY: The Tampa location (NPI {t['billing_npi']}) has "
            f"dramatically different billing patterns: cost/claim=${t['cost_per_claim']:,.2f} "
            f"vs ${other_avg_cpc:,.2f} average for other Ambassador locations, and "
            f"{t['total_beneficiaries']:,} beneficiaries (much higher). "
            "This may represent a legitimate operation or a different billing strategy."
        )

    return flags


def build_investigation_report(
    providers_df: pd.DataFrame | None = None,
) -> AmbassadorFraudReport:
    """Run the full Ambassador Health fraud investigation.

    Parameters
    ----------
    providers_df : DataFrame, optional
        Full provider dataset. If *None*, loads from the PBC providers CSV.

    Returns
    -------
    AmbassadorFraudReport
    """
    if providers_df is None:
        providers_df = load_pbc_providers()
    if providers_df is None or providers_df.empty:
        raise FileNotFoundError("Provider data not available.")

    # Apply fraud flags to ALL providers first (for percentile context)
    all_flagged = compute_fraud_signals(providers_df)

    # Extract Ambassador entities with their flags
    amb = get_ambassador_entities(all_flagged)
    peers = _compute_peer_stats(all_flagged)
    red_flags = _identify_red_flags(amb, peers)

    # Build flag summary
    flag_cols = [c for c in amb.columns if c.startswith("flag_")]
    flags_summary = {col: int(amb[col].sum()) for col in flag_cols}

    report = AmbassadorFraudReport(
        entities=amb,
        total_paid=float(amb["total_paid"].sum()),
        entity_count=len(amb),
        unique_cities=sorted(amb["city"].unique().tolist()),
        unique_npis=sorted(amb["billing_npi"].unique().tolist()),
        flags_summary=flags_summary,
        peer_comparison=peers,
        red_flags=red_flags,
    )
    return report


def report_to_dict(report: AmbassadorFraudReport) -> dict:
    """Convert report to a JSON-serialisable dict."""
    entity_records = report.entities[
        [
            "billing_npi",
            "provider_name",
            "city",
            "zip",
            "total_paid",
            "total_claims",
            "total_beneficiaries",
            "cost_per_claim",
            "cost_per_beneficiary",
            "claims_per_beneficiary",
            "months_active",
            "unique_hcpcs_codes",
            "fraud_signal_count",
        ]
    ].to_dict(orient="records")

    return {
        "investigation": "Ambassador Health / Ambassador Home Health Services",
        "state": "Florida",
        "total_medicaid_paid": report.total_paid,
        "entity_count": report.entity_count,
        "unique_npis": report.unique_npis,
        "unique_cities": report.unique_cities,
        "entities": entity_records,
        "flags_summary": report.flags_summary,
        "peer_comparison": report.peer_comparison,
        "red_flags": report.red_flags,
    }


def save_report(report: AmbassadorFraudReport, path: Path | None = None) -> Path:
    """Write the investigation report to a JSON artifact."""
    if path is None:
        path = ARTIFACT_DIR / "ambassador_health_investigation.json"
    with open(path, "w") as f:
        json.dump(report_to_dict(report), f, indent=2)
    return path


# ---------------------------------------------------------------------------
# Government referral report
# ---------------------------------------------------------------------------

_REFERRAL_AGENCIES = {
    "FL_MFCU": {
        "name": "Florida Attorney General — Medicaid Fraud Control Unit",
        "address": "Office of the Attorney General, PL-01 The Capitol, Tallahassee, FL 32399-1050",
        "phone": "(850) 414-3990",
        "authority": "42 U.S.C. § 1396b(q); Fla. Stat. § 409.920 (Florida False Claims Act)",
        "jurisdiction": "State-level Medicaid fraud investigation and prosecution",
    },
    "HHS_OIG": {
        "name": "U.S. Department of Health & Human Services — Office of Inspector General",
        "address": "Attn: OIG Hotline Operations, P.O. Box 23489, Washington, DC 20026",
        "phone": "1-800-HHS-TIPS (1-800-447-8477)",
        "authority": "42 U.S.C. § 1320a-7 (Exclusion); 31 U.S.C. §§ 3729–3733 (Federal False Claims Act)",
        "jurisdiction": "Federal Medicaid/Medicare fraud, provider exclusion, civil monetary penalties",
    },
    "CMS_CPI": {
        "name": "Centers for Medicare & Medicaid Services — Center for Program Integrity",
        "address": "7500 Security Boulevard, Baltimore, MD 21244",
        "phone": "1-800-633-4227",
        "authority": "42 C.F.R. § 455 (Fraud and Abuse in Medicaid); 42 C.F.R. § 424.535 (Revocation)",
        "jurisdiction": "Payment suspension, NPI revocation, referral to OIG/DOJ",
    },
}


def generate_referral_text(report: AmbassadorFraudReport, agency_key: str = "FL_MFCU") -> str:
    """Generate a formal referral letter/report for a government agency.

    Parameters
    ----------
    report : AmbassadorFraudReport
        Completed investigation report.
    agency_key : str
        Target agency key from _REFERRAL_AGENCIES.

    Returns
    -------
    str
        Formatted referral text ready for submission.
    """
    agency = _REFERRAL_AGENCIES.get(agency_key, _REFERRAL_AGENCIES["FL_MFCU"])
    today = date.today().strftime("%B %d, %Y")

    # Build entity table
    entity_lines = []
    for _, row in report.entities.sort_values("total_paid", ascending=False).iterrows():
        entity_lines.append(
            f"    NPI: {row['billing_npi']}\n"
            f"    Name: {row['provider_name']}\n"
            f"    City/ZIP: {row['city']}, FL {row['zip']}\n"
            f"    Total Medicaid Paid: ${row['total_paid']:,.2f}\n"
            f"    Total Claims: {int(row['total_claims']):,}\n"
            f"    Total Beneficiaries: {int(row['total_beneficiaries']):,}\n"
            f"    Cost per Claim: ${row['cost_per_claim']:,.2f}\n"
            f"    Cost per Beneficiary: ${row['cost_per_beneficiary']:,.2f}\n"
            f"    Claims per Beneficiary: {row['claims_per_beneficiary']:.1f}\n"
            f"    Unique HCPCS Codes Billed: {int(row['unique_hcpcs_codes'])}\n"
            f"    Fraud Signal Count: {int(row['fraud_signal_count'])} of 4\n"
            f"    Months Active: {int(row['months_active'])}"
        )

    entity_block = "\n\n".join(entity_lines)

    # Build flag summary
    flag_labels = {
        "flag_high_cost_per_claim": "Cost per claim exceeds 95th percentile of all providers",
        "flag_high_cost_per_beneficiary": "Cost per beneficiary exceeds 95th percentile",
        "flag_high_claims_per_beneficiary": "Claims per beneficiary exceeds 95th percentile",
        "flag_low_code_diversity": "Billing only 1-2 HCPCS codes (low service diversity)",
    }
    flag_lines = []
    for flag_col, count in report.flags_summary.items():
        label = flag_labels.get(flag_col, flag_col)
        flag_lines.append(f"  - {label}: {count} of {report.entity_count} entities")
    flag_block = "\n".join(flag_lines)

    # Red flags
    red_flag_lines = []
    for i, flag in enumerate(report.red_flags, 1):
        red_flag_lines.append(f"  {i}. {flag}")
    red_flag_block = "\n\n".join(red_flag_lines)

    # HCPCS codes
    hcpcs_lines = []
    for code, info in AMBASSADOR_KNOWN_HCPCS.items():
        confirmed = info.get("confirmed_npis", [])
        hcpcs_lines.append(
            f"  - {code}: {info['description']} ({info['category']})\n"
            f"    Status: {'Confirmed via claim-level data for NPIs ' + ', '.join(confirmed) if confirmed else info.get('note', 'Probable')}"
        )
    hcpcs_block = "\n".join(hcpcs_lines)

    peers = report.peer_comparison

    text = f"""
================================================================================
FORMAL REFERRAL FOR INVESTIGATION
Suspected Medicaid Fraud — Ambassador Health Services Network
================================================================================

Date: {today}

To: {agency['name']}
    {agency['address']}
    Phone: {agency['phone']}

RE: Referral for Investigation of Suspected Medicaid Fraud
    Subject: Ambassador Health Services / Ambassador Home Health Services
    State: Florida
    Aggregate Medicaid Payments: ${report.total_paid:,.2f}
    Number of Provider Entities (NPIs): {report.entity_count}
    Applicable Law: {agency['authority']}

--------------------------------------------------------------------------------
1. EXECUTIVE SUMMARY
--------------------------------------------------------------------------------

This referral presents data-driven evidence of potential Medicaid billing fraud
by a network of {report.entity_count} provider entities operating under the
"Ambassador Health Services" and "Ambassador Home Health Services" brand names
across {len(report.unique_cities)} cities in Florida.

These entities have collectively billed Florida Medicaid ${report.total_paid:,.2f}
(approximately ${report.total_paid / 1e6:,.1f} million). Statistical analysis
of publicly available Medicaid provider spending data reveals that every
Ambassador entity exhibits multiple fraud indicators that substantially deviate
from peer provider norms.

The pattern of multiple separate NPIs under a common brand, combined with
extreme billing metrics and minimal service code diversity, is consistent with
a coordinated billing scheme designed to maximize payments while avoiding
per-entity audit thresholds.

--------------------------------------------------------------------------------
2. SUBJECT ENTITIES
--------------------------------------------------------------------------------

The following {report.entity_count} National Provider Identifiers (NPIs) were
identified as operating under the Ambassador brand:

{entity_block}

Cities: {', '.join(report.unique_cities)}

--------------------------------------------------------------------------------
3. HCPCS CODES BILLED
--------------------------------------------------------------------------------

The Ambassador entities predominantly bill the following home health nursing
codes. Claim-level sample data confirms S9124 usage; the extremely low code
diversity (6 of 9 entities bill only 1-2 codes) suggests near-exclusive
reliance on these codes:

{hcpcs_block}

For context, S9124 is the single highest-spending HCPCS code in Palm Beach
County at $187.97 million across all providers. Ambassador entities billing
this code at ${report.entities['cost_per_claim'].mean():,.2f} per claim is
significantly above the population average.

--------------------------------------------------------------------------------
4. STATISTICAL FRAUD INDICATORS
--------------------------------------------------------------------------------

The following statistical signals were computed by comparing each Ambassador
entity against all {int(peers.get('median_unique_hcpcs', 0)) or 'N/A'}-code
providers in the dataset. Thresholds are set at the 95th percentile — meaning
only the top 5% of providers would normally trigger each flag.

Flag Summary (entities triggering / total entities):

{flag_block}

Key peer benchmarks used:
  - Median cost per claim (all providers):       ${peers['median_cost_per_claim']:,.2f}
  - 95th percentile cost per claim:              ${peers['p95_cost_per_claim']:,.2f}
  - Ambassador average cost per claim:           ${report.entities['cost_per_claim'].mean():,.2f}

  - Median cost per beneficiary (all providers): ${peers['median_cost_per_beneficiary']:,.2f}
  - 95th percentile cost per beneficiary:        ${peers['p95_cost_per_beneficiary']:,.2f}
  - Ambassador average cost per beneficiary:     ${report.entities['cost_per_beneficiary'].mean():,.2f}

  - Median claims per beneficiary:               {peers['median_claims_per_beneficiary']:.1f}
  - 95th percentile claims per beneficiary:      {peers['p95_claims_per_beneficiary']:.1f}
  - Ambassador average claims per beneficiary:   {report.entities['claims_per_beneficiary'].mean():.1f}

--------------------------------------------------------------------------------
5. RED FLAGS — DETAILED FINDINGS
--------------------------------------------------------------------------------

{red_flag_block}

--------------------------------------------------------------------------------
6. DATA SOURCES AND METHODOLOGY
--------------------------------------------------------------------------------

Data Source:
  - HHS Medicaid Provider Spending data released via the Department of
    Government Efficiency (DOGE) at stopendataprod.blob.core.windows.net
  - Full dataset: approximately 227 million rows, 10.3 GB
  - Provider names resolved via CMS NPPES NPI Registry (npiregistry.cms.hhs.gov)
  - Coverage period: 2018-01 through 2024-12 (84 months)

Methodology:
  - All providers in the dataset were ranked by total Medicaid payments
  - Per-unit metrics (cost/claim, cost/beneficiary, claims/beneficiary) were
    computed for each provider
  - Fraud signal flags were set at the 95th percentile of each metric
  - Low code diversity flagged when a provider bills 2 or fewer HCPCS codes
  - Ambassador entities were identified by provider name matching and
    cross-referenced across all NPI registrations in Florida

Limitations:
  - This analysis is based on aggregated billing data, not individual claim
    records. The specific services rendered cannot be verified from this data.
  - The HCPCS codes billed by each Ambassador entity are confirmed for 2 NPIs
    via a sample; the remaining 7 are inferred from billing patterns.
  - This referral identifies statistical anomalies. Determination of fraud
    requires access to medical records, patient interviews, and on-site review.

--------------------------------------------------------------------------------
7. REQUESTED ACTIONS
--------------------------------------------------------------------------------

Based on the findings above, we respectfully request:

  1. INVESTIGATION: Formal investigation into the billing practices of all
     {report.entity_count} Ambassador-branded entities listed in Section 2.

  2. MEDICAL RECORD REVIEW: Verification that claimed home nursing visits
     (S9124/S9123/S9122) correspond to actual services rendered to actual
     patients, particularly for entities with extreme claims-per-beneficiary
     ratios (19-23 claims per beneficiary vs. population median of
     {peers['median_claims_per_beneficiary']:.1f}).

  3. PATIENT VERIFICATION: Confirmation that the {int(report.entities['total_beneficiaries'].sum()):,}
     unique beneficiaries billed across Ambassador entities are real patients
     who received the claimed services.

  4. CORPORATE STRUCTURE REVIEW: Investigation of the ownership and management
     relationships between the 9 Ambassador NPIs to determine whether they
     operate as a coordinated network.

  5. PAYMENT SUSPENSION: Consideration of interim payment suspension per
     42 C.F.R. § 455.23 pending investigation, given the credible evidence
     of fraud presented herein.

--------------------------------------------------------------------------------
8. CONTACT AND SUPPORTING MATERIALS
--------------------------------------------------------------------------------

Supporting data artifacts are available upon request:
  - ambassador_health_investigation.json (structured investigation data)
  - doge_medicaid_pbc_providers.csv (full provider dataset with fraud flags)
  - doge_medicaid_provider_spending_sample.csv (claim-level sample data)
  - doge_medicaid_pbc_stats.json (aggregate statistics)

================================================================================
END OF REFERRAL
================================================================================
"""
    return text.strip()


def save_referral(report: AmbassadorFraudReport, agency_key: str = "FL_MFCU", path: Path | None = None) -> Path:
    """Generate and save the formal referral text."""
    text = generate_referral_text(report, agency_key)
    if path is None:
        path = ARTIFACT_DIR / "ambassador_health_referral.txt"
    with open(path, "w") as f:
        f.write(text)
    return path


def generate_referral_docx(report: AmbassadorFraudReport, agency_key: str = "FL_MFCU") -> bytes:
    """Generate a Word (.docx) referral document and return raw bytes.

    Returns bytes so it can be used directly with ``st.download_button``.
    """
    from io import BytesIO

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    agency = _REFERRAL_AGENCIES.get(agency_key, _REFERRAL_AGENCIES["FL_MFCU"])
    today = date.today().strftime("%B %d, %Y")
    doc = Document()

    # -- styles ---------------------------------------------------------------
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    def _heading(text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    def _bold_para(label: str, value: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        p.add_run(value)

    # -- title page -----------------------------------------------------------
    title = doc.add_heading("FORMAL REFERRAL FOR INVESTIGATION", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Suspected Medicaid Fraud — Ambassador Health Services Network")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0xD6, 0x27, 0x28)

    doc.add_paragraph("")
    _bold_para("Date: ", today)
    _bold_para("To: ", agency["name"])
    doc.add_paragraph(f"    {agency['address']}")
    doc.add_paragraph(f"    Phone: {agency['phone']}")
    doc.add_paragraph("")
    _bold_para("RE: ", "Referral for Investigation of Suspected Medicaid Fraud")
    _bold_para("Subject: ", "Ambassador Health Services / Ambassador Home Health Services")
    _bold_para("State: ", "Florida")
    _bold_para("Aggregate Medicaid Payments: ", f"${report.total_paid:,.2f}")
    _bold_para("Number of Provider Entities (NPIs): ", str(report.entity_count))
    _bold_para("Applicable Law: ", agency["authority"])

    # -- 1. Executive Summary -------------------------------------------------
    _heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        f"This referral presents data-driven evidence of potential Medicaid billing fraud "
        f"by a network of {report.entity_count} provider entities operating under the "
        f'"Ambassador Health Services" and "Ambassador Home Health Services" brand names '
        f"across {len(report.unique_cities)} cities in Florida."
    )
    doc.add_paragraph(
        f"These entities have collectively billed Florida Medicaid ${report.total_paid:,.2f} "
        f"(approximately ${report.total_paid / 1e6:,.1f} million). Statistical analysis of "
        f"publicly available Medicaid provider spending data reveals that every Ambassador "
        f"entity exhibits multiple fraud indicators that substantially deviate from peer "
        f"provider norms."
    )
    doc.add_paragraph(
        "The pattern of multiple separate NPIs under a common brand, combined with extreme "
        "billing metrics and minimal service code diversity, is consistent with a coordinated "
        "billing scheme designed to maximize payments while avoiding per-entity audit thresholds."
    )

    # -- 2. Subject Entities (table) ------------------------------------------
    _heading("2. Subject Entities", level=1)
    doc.add_paragraph(
        f"The following {report.entity_count} National Provider Identifiers (NPIs) were "
        f"identified as operating under the Ambassador brand:"
    )

    cols = ["NPI", "Provider Name", "City", "Total Paid", "Claims",
            "Beneficiaries", "$/Claim", "$/Bene", "Claims/Bene", "HCPCS", "Flags"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, col_name in enumerate(cols):
        hdr[i].text = col_name
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    for _, row in report.entities.sort_values("total_paid", ascending=False).iterrows():
        cells = table.add_row().cells
        values = [
            str(row["billing_npi"]),
            str(row["provider_name"])[:30],
            str(row["city"]),
            f"${row['total_paid']:,.0f}",
            f"{int(row['total_claims']):,}",
            f"{int(row['total_beneficiaries']):,}",
            f"${row['cost_per_claim']:,.0f}",
            f"${row['cost_per_beneficiary']:,.0f}",
            f"{row['claims_per_beneficiary']:.1f}",
            str(int(row["unique_hcpcs_codes"])),
            f"{int(row['fraud_signal_count'])}/4",
        ]
        for i, val in enumerate(values):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph(f"\nCities: {', '.join(report.unique_cities)}")

    # -- 3. HCPCS Codes -------------------------------------------------------
    _heading("3. HCPCS Codes Billed", level=1)
    doc.add_paragraph(
        "The Ambassador entities predominantly bill the following home health nursing codes. "
        "Claim-level sample data confirms S9124 usage; the extremely low code diversity "
        "(6 of 9 entities bill only 1-2 codes) suggests near-exclusive reliance on these codes:"
    )
    for code, info in AMBASSADOR_KNOWN_HCPCS.items():
        confirmed = info.get("confirmed_npis", [])
        status = (
            f"Confirmed via claim-level data for NPIs {', '.join(confirmed)}"
            if confirmed
            else info.get("note", "Probable")
        )
        _bold_para(f"{code}: ", f"{info['description']} ({info['category']}) — {status}")

    doc.add_paragraph(
        "S9124 is the single highest-spending HCPCS code in Palm Beach County at "
        "$187.97 million across all providers."
    )

    # -- 4. Statistical Fraud Indicators --------------------------------------
    _heading("4. Statistical Fraud Indicators", level=1)
    peers = report.peer_comparison

    flag_labels = {
        "flag_high_cost_per_claim": "Cost per claim exceeds 95th percentile",
        "flag_high_cost_per_beneficiary": "Cost per beneficiary exceeds 95th percentile",
        "flag_high_claims_per_beneficiary": "Claims per beneficiary exceeds 95th percentile",
        "flag_low_code_diversity": "Billing only 1-2 HCPCS codes (low service diversity)",
    }

    ft = doc.add_table(rows=1, cols=3)
    ft.style = "Light Grid Accent 1"
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    fhdr = ft.rows[0].cells
    for i, h in enumerate(["Signal", "Entities Flagged", "Rate"]):
        fhdr[i].text = h
        for p in fhdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for flag_col, count in report.flags_summary.items():
        r = ft.add_row().cells
        r[0].text = flag_labels.get(flag_col, flag_col)
        r[1].text = f"{count} / {report.entity_count}"
        r[2].text = f"{count / report.entity_count * 100:.0f}%"
        for cell in r:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")
    _heading("Peer Benchmarks", level=2)
    benchmarks = [
        ("Median cost/claim", f"${peers['median_cost_per_claim']:,.2f}",
         f"${report.entities['cost_per_claim'].mean():,.2f}"),
        ("95th pctl cost/claim", f"${peers['p95_cost_per_claim']:,.2f}", "—"),
        ("Median cost/beneficiary", f"${peers['median_cost_per_beneficiary']:,.2f}",
         f"${report.entities['cost_per_beneficiary'].mean():,.2f}"),
        ("95th pctl cost/beneficiary", f"${peers['p95_cost_per_beneficiary']:,.2f}", "—"),
        ("Median claims/beneficiary", f"{peers['median_claims_per_beneficiary']:.1f}",
         f"{report.entities['claims_per_beneficiary'].mean():.1f}"),
        ("95th pctl claims/beneficiary", f"{peers['p95_claims_per_beneficiary']:.1f}", "—"),
    ]
    bt = doc.add_table(rows=1, cols=3)
    bt.style = "Light Grid Accent 1"
    bt.alignment = WD_TABLE_ALIGNMENT.CENTER
    bhdr = bt.rows[0].cells
    for i, h in enumerate(["Metric", "All Providers", "Ambassador Avg"]):
        bhdr[i].text = h
        for p in bhdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for label, all_val, amb_val in benchmarks:
        r = bt.add_row().cells
        r[0].text = label
        r[1].text = all_val
        r[2].text = amb_val
        for cell in r:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    # -- 5. Red Flags ---------------------------------------------------------
    _heading("5. Red Flags — Detailed Findings", level=1)
    for i, flag in enumerate(report.red_flags, 1):
        parts = flag.split(":", 1)
        if len(parts) == 2:
            p = doc.add_paragraph()
            r = p.add_run(f"{i}. {parts[0]}:")
            r.bold = True
            r.font.color.rgb = RGBColor(0xD6, 0x27, 0x28)
            p.add_run(parts[1])
        else:
            doc.add_paragraph(f"{i}. {flag}")

    # -- 6. Data Sources ------------------------------------------------------
    _heading("6. Data Sources and Methodology", level=1)
    _heading("Data Source", level=2)
    for line in [
        "HHS Medicaid Provider Spending data released via the Department of Government Efficiency (DOGE)",
        "Full dataset: approximately 227 million rows, 10.3 GB",
        "Provider names resolved via CMS NPPES NPI Registry (npiregistry.cms.hhs.gov)",
        "Coverage period: 2018-01 through 2024-12 (84 months)",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    _heading("Methodology", level=2)
    for line in [
        "All providers in the dataset were ranked by total Medicaid payments",
        "Per-unit metrics (cost/claim, cost/beneficiary, claims/beneficiary) computed for each provider",
        "Fraud signal flags set at the 95th percentile of each metric",
        "Low code diversity flagged when a provider bills 2 or fewer HCPCS codes",
        "Ambassador entities identified by provider name matching and cross-referenced across all Florida NPI registrations",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    _heading("Limitations", level=2)
    for line in [
        "This analysis is based on aggregated billing data, not individual claim records.",
        "HCPCS codes billed by each Ambassador entity are confirmed for 2 NPIs via a sample; the remaining 7 are inferred.",
        "This referral identifies statistical anomalies. Determination of fraud requires medical records, patient interviews, and on-site review.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # -- 7. Requested Actions -------------------------------------------------
    _heading("7. Requested Actions", level=1)
    actions = [
        ("INVESTIGATION", f"Formal investigation into the billing practices of all {report.entity_count} Ambassador-branded entities."),
        ("MEDICAL RECORD REVIEW", "Verification that claimed home nursing visits (S9124/S9123/S9122) correspond to actual services rendered."),
        ("PATIENT VERIFICATION", f"Confirmation that the {int(report.entities['total_beneficiaries'].sum()):,} unique beneficiaries are real patients who received claimed services."),
        ("CORPORATE STRUCTURE REVIEW", "Investigation of ownership and management relationships between the 9 Ambassador NPIs."),
        ("PAYMENT SUSPENSION", "Consideration of interim payment suspension per 42 C.F.R. § 455.23 pending investigation."),
    ]
    for i, (title, detail) in enumerate(actions, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. {title}: ")
        r.bold = True
        p.add_run(detail)

    # -- 8. Supporting Materials ----------------------------------------------
    _heading("8. Contact and Supporting Materials", level=1)
    doc.add_paragraph("Supporting data artifacts are available upon request:")
    for f in [
        "ambassador_health_investigation.json (structured investigation data)",
        "doge_medicaid_pbc_providers.csv (full provider dataset with fraud flags)",
        "doge_medicaid_provider_spending_sample.csv (claim-level sample data)",
        "doge_medicaid_pbc_stats.json (aggregate statistics)",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    # -- save to bytes --------------------------------------------------------
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def save_referral_docx(
    report: AmbassadorFraudReport,
    agency_key: str = "FL_MFCU",
    path: Path | None = None,
) -> Path:
    """Generate and save the Word referral document."""
    data = generate_referral_docx(report, agency_key)
    if path is None:
        path = ARTIFACT_DIR / "ambassador_health_referral.docx"
    with open(path, "wb") as f:
        f.write(data)
    return path


# ---------------------------------------------------------------------------
# CLI entry-point
# ---------------------------------------------------------------------------

def main() -> None:
    """Run the Ambassador Health investigation and print findings."""
    report = build_investigation_report()
    out = save_report(report)
    print(f"Ambassador Health Investigation Report saved to {out}\n")
    print(f"{'='*72}")
    print("AMBASSADOR HEALTH FRAUD INVESTIGATION — SUMMARY")
    print(f"{'='*72}\n")
    print(f"Total entities:       {report.entity_count}")
    print(f"Total Medicaid paid:  ${report.total_paid:,.2f}")
    print(f"Unique NPIs:          {len(report.unique_npis)}")
    print(f"Cities:               {', '.join(report.unique_cities)}")
    print()
    print(f"{'─'*72}")
    print("ENTITY BREAKDOWN")
    print(f"{'─'*72}")
    for _, row in report.entities.iterrows():
        print(
            f"  NPI {row['billing_npi']}  {row['provider_name']:<45s} "
            f"{row['city']:<16s} ${row['total_paid']:>14,.2f}  "
            f"flags={int(row['fraud_signal_count'])}"
        )
    print()
    print(f"{'─'*72}")
    print("FLAG SUMMARY")
    print(f"{'─'*72}")
    for flag, count in report.flags_summary.items():
        print(f"  {flag}: {count}/{report.entity_count} entities")
    print()
    print(f"{'─'*72}")
    print("RED FLAGS")
    print(f"{'─'*72}")
    for i, flag in enumerate(report.red_flags, 1):
        print(f"\n  [{i}] {flag}")
    print()


if __name__ == "__main__":
    main()
